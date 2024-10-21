from docx import Document
from docx.shared import Inches

# 문서 생성
doc = Document()

# 제목 추가
doc.add_heading('자판기 프로그램 명세서', 0)

# 개요 추가
doc.add_heading('개요', level=1)
doc.add_paragraph(
    '이 프로그램은 자판기의 상태를 관리하고, 동전 투입, 상품 선택, 상품 수령, 거스름돈 반환 등의 기능을 제공하는 자판기 시뮬레이션입니다. '
    '상태 패턴을 사용하여 자판기의 다양한 상태(초기 상태, 금액 수령 상태, 상품 제공 상태, 거스름돈 반환 상태, 반환 처리 상태)를 관리합니다.'
)

# 클래스 목록 및 메서드 목록 추가
doc.add_heading('클래스 목록 및 메서드 목록', level=1)

# State 클래스
doc.add_heading('State 클래스', level=2)
doc.add_paragraph('메서드:')
doc.add_paragraph('insert_coin(vending_machine, amount)', style='List Bullet')
doc.add_paragraph('select_product(vending_machine, product)', style='List Bullet')
doc.add_paragraph('receive_product(vending_machine)', style='List Bullet')
doc.add_paragraph('receive_change(vending_machine)', style='List Bullet')
doc.add_paragraph('press_return_button(vending_machine)', style='List Bullet')
doc.add_paragraph('receive_returned_amount(vending_machine)', style='List Bullet')

# ReceivingAmountState 클래스
doc.add_heading('ReceivingAmountState 클래스', level=2)
doc.add_paragraph('메서드:')
doc.add_paragraph('insert_coin(vending_machine, amount)', style='List Bullet')
doc.add_paragraph('select_product(vending_machine, product)', style='List Bullet')

# 필드 및 메서드 설명 추가
doc.add_heading('필드 및 메서드 설명', level=1)

# State 클래스 설명
doc.add_heading('State 클래스', level=2)
doc.add_paragraph(
    'State 클래스는 자판기의 상태를 나타내는 추상 클래스입니다. '
    '다음과 같은 추상 메서드를 포함합니다:'
)
doc.add_paragraph('insert_coin(vending_machine, amount): 동전을 투입하는 메서드입니다.', style='List Bullet')
doc.add_paragraph('select_product(vending_machine, product): 상품을 선택하는 메서드입니다.', style='List Bullet')
doc.add_paragraph('receive_product(vending_machine): 상품을 수령하는 메서드입니다.', style='List Bullet')
doc.add_paragraph('receive_change(vending_machine): 거스름돈을 받는 메서드입니다.', style='List Bullet')
doc.add_paragraph('press_return_button(vending_machine): 반환 버튼을 누르는 메서드입니다.', style='List Bullet')
doc.add_paragraph('receive_returned_amount(vending_machine): 반환된 금액을 받는 메서드입니다.', style='List Bullet')

# ReceivingAmountState 클래스 설명
doc.add_heading('ReceivingAmountState 클래스', level=2)
doc.add_paragraph(
    'ReceivingAmountState 클래스는 State 클래스를 상속받아 금액 수령 상태를 구현합니다. '
    '다음과 같은 메서드를 포함합니다:'
)

# insert_coin 메서드 설명
doc.add_heading('insert_coin 메서드', level=3)
doc.add_paragraph('동전을 투입하여 금액을 추가합니다.')
table = doc.add_table(rows=3, cols=2)
table.style = 'Table Grid'
table.cell(0, 0).text = '매개변수'
table.cell(0, 1).text = '설명'
table.cell(1, 0).text = 'vending_machine'
table.cell(1, 1).text = '자판기 객체'
table.cell(2, 0).text = 'amount'
table.cell(2, 1).text = '투입된 금액'
doc.add_paragraph('내부 동작:')
doc.add_paragraph('- 투입된 금액을 자판기의 총 금액에 추가합니다.', style='List Bullet')
doc.add_paragraph('- 현재 합계 금액을 출력합니다.', style='List Bullet')

# select_product 메서드 설명
doc.add_heading('select_product 메서드', level=3)
doc.add_paragraph('상품을 선택하여 금액과 재고를 확인하고, 조건에 따라 상품을 제공하거나 오류 메시지를 출력합니다.')
table = doc.add_table(rows=3, cols=2)
table.style = 'Table Grid'
table.cell(0, 0).text = '매개변수'
table.cell(0, 1).text = '설명'
table.cell(1, 0).text = 'vending_machine'
table.cell(1, 1).text = '자판기 객체'
table.cell(2, 0).text = 'product'
table.cell(2, 1).text = '선택된 상품'
doc.add_paragraph('내부 동작:')
doc.add_paragraph('- 선택된 상품이 자판기 재고에 있는지 확인합니다.', style='List Bullet')
doc.add_paragraph('- 상품의 가격을 확인합니다.', style='List Bullet')
doc.add_paragraph('- 투입된 금액이 상품의 가격보다 크거나 같은지 확인합니다.', style='List Bullet')
doc.add_paragraph('- 재고가 있는 경우 상품을 제공하고, 재고를 감소시킵니다.', style='List Bullet')
doc.add_paragraph('- 재고가 없는 경우 오류 메시지를 출력합니다.', style='List Bullet')
doc.add_paragraph('- 금액이 부족한 경우 오류 메시지를 출력합니다.', style='List Bullet')

# press_return_button 메서드 설명
doc.add_heading('press_return_button 메서드', level=3)
doc.add_paragraph('반환 버튼을 누르는 메서드입니다.')
table = doc.add_table(rows=2, cols=2)
table.style = 'Table Grid'
table.cell(0, 0).text = '매개변수'
table.cell(0, 1).text = '설명'
table.cell(1, 0).text = 'vending_machine'
table.cell(1, 1).text = '자판기 객체'
doc.add_paragraph('내부 동작:')
doc.add_paragraph('- 동전을 먼저 넣어달라는 메시지를 출력합니다.', style='List Bullet')

# receive_returned_amount 메서드 설명
doc.add_heading('receive_returned_amount 메서드', level=3)
doc.add_paragraph('반환된 금액을 받는 메서드입니다.')
table = doc.add_table(rows=2, cols=2)
table.style = 'Table Grid'
table.cell(0, 0).text = '매개변수'
table.cell(0, 1).text = '설명'
table.cell(1, 0).text = 'vending_machine'
table.cell(1, 1).text = '자판기 객체'
doc.add_paragraph('내부 동작:')
doc.add_paragraph('- 동전을 먼저 넣어달라는 메시지를 출력합니다.', style='List Bullet')

# 문서 저장
doc.save('VendingMachine_Specification.docx')